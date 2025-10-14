"""
Update DOCX templates to include company address placeholder.
"""
from docx import Document
from docx.shared import Pt

def update_template(template_path, language='en'):
    """Add company address placeholder to template."""
    doc = Document(template_path)
    
    # The structure should be:
    # 0: (empty)
    # 1: {{DATE}}
    # 2: (new) {{COMPANY_ADDRESS_LINE1}}
    # 3: (new) {{COMPANY_ADDRESS_LINE2}}
    # 4: (new empty line)
    # 5: Application as {{JOB_TITLE}} / Bewerbung als {{JOB_TITLE}}
    # ...
    
    # Insert 3 new paragraphs after the date (paragraph 1)
    # We'll insert them in reverse order to maintain correct indexing
    
    # Get the paragraph after which we want to insert (the date paragraph)
    date_para = doc.paragraphs[1]
    
    # Insert empty line
    p_empty = date_para.insert_paragraph_before()
    
    # Insert address line 2
    p_addr2 = date_para.insert_paragraph_before()
    p_addr2.text = '{{COMPANY_ADDRESS_LINE2}}'
    p_addr2.style = date_para.style
    for run in p_addr2.runs:
        run.font.size = Pt(11)
    
    # Insert address line 1
    p_addr1 = date_para.insert_paragraph_before()
    p_addr1.text = '{{COMPANY_ADDRESS_LINE1}}'
    p_addr1.style = date_para.style
    for run in p_addr1.runs:
        run.font.size = Pt(11)
    
    # Insert empty line before address
    p_empty_before = date_para.insert_paragraph_before()
    
    # Actually, let's do this differently - we'll work with the current structure
    # and insert after paragraph 1 (date)
    
    # Clear what we just did and start over
    doc = Document(template_path)
    
    # Simpler approach: manually add after date
    paragraphs = list(doc.paragraphs)
    
    # Find the date paragraph
    for i, para in enumerate(paragraphs):
        if '{{DATE}}' in para.text:
            print(f"Found {{{{DATE}}}} at paragraph {i}")
            # We need to insert new paragraphs after this one
            # But python-docx doesn't have an easy "insert_paragraph_after" method
            # So let's use a different approach: modify the XML directly
            
            # Actually, let's just modify the existing template manually
            # or use the _element approach
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Get the paragraph element
            p_element = para._element
            
            # Create new paragraph elements
            new_p1 = OxmlElement('w:p')
            new_p2 = OxmlElement('w:p')
            new_p3 = OxmlElement('w:p')
            new_p4 = OxmlElement('w:p')
            
            # Add runs with text to new paragraphs
            for new_p, text in [(new_p1, ''), (new_p2, '{{COMPANY_ADDRESS_LINE1}}'), 
                                (new_p3, '{{COMPANY_ADDRESS_LINE2}}'), (new_p4, '')]:
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = text
                r.append(t)
                new_p.append(r)
            
            # Insert after the date paragraph
            parent = p_element.getparent()
            parent.insert(parent.index(p_element) + 1, new_p1)
            parent.insert(parent.index(p_element) + 2, new_p2)
            parent.insert(parent.index(p_element) + 3, new_p3)
            parent.insert(parent.index(p_element) + 4, new_p4)
            
            break
    
    # Save the modified template
    backup_path = template_path.replace('.docx', '_backup.docx')
    import shutil
    shutil.copy(template_path, backup_path)
    print(f"Backup created: {backup_path}")
    
    doc.save(template_path)
    print(f"Template updated: {template_path}")
    
    # Verify
    doc_verify = Document(template_path)
    print(f"\nUpdated template structure:")
    for i, p in enumerate(doc_verify.paragraphs[:10]):
        print(f"{i}: {p.text}")

if __name__ == '__main__':
    print("Updating English template...")
    update_template('data/template_en.docx', 'en')
    
    print("\n" + "="*80 + "\n")
    
    print("Updating German template...")
    update_template('data/template_de.docx', 'de')
