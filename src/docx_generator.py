"""
Word Template-based Cover Letter Generator
Uses .docx templates with placeholders for easy customization
"""


from docx import Document
from datetime import datetime
import os
import re
from typing import Dict, Any, Optional
try:
    from .utils.log_config import get_logger
    from .utils.errors import DocumentError
except Exception:
    import sys
    sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from utils.log_config import get_logger
    from utils.errors import DocumentError


class WordCoverLetterGenerator:
    """
    Generates cover letters from Word templates
    """
    
    def __init__(self, template_path_de: str = 'data/template_de.docx', template_path_en: str = 'data/template_en.docx') -> None:
        self.template_de = template_path_de
        self.template_en = template_path_en
        
        # Your personal details
        self.sender = {
            'name': 'Dr. Kai Voges',
            'phone': '+49 1622146092',
            'email': 'kai.voges@gmx.net',
            'address_line1': 'Cranachstraße 4',
            'address_line2': '40235 Düsseldorf',
            'linkedin': 'www.linkedin.com/in/worldapprentice'
        }
    
    def generate_from_template(
        self,
        cover_letter_text: str,
        job_data: Dict[str, Any],
        output_path: str,
        language: str = 'german'
    ) -> str:
        """
        Generate cover letter from Word template
        
        Args:
            cover_letter_text (str): The AI-generated cover letter content
            job_data (dict): Job information
            output_path (str): Where to save the output
            language (str): 'german' or 'english'
            
        Returns:
            str: Path to generated document
        """
        logger = get_logger(__name__)
        logger.info("Generating Word Document | lang=%s | out=%s", language, output_path)

        # Select template
        template_path = self.template_de if language == 'german' else self.template_en

        # Check if template exists
        if not os.path.exists(template_path):
            logger.warning("Template not found: %s. Falling back to basic generation.", template_path)
            return self._generate_basic_docx(cover_letter_text, job_data, output_path, language)

        # Load template
        try:
            doc = Document(template_path)
        except Exception as e:
            logger.error("Failed to load template %s: %s", template_path, e)
            raise DocumentError(f"Failed to load template: {template_path}") from e

        # Prepare replacement data
        today = datetime.now().strftime('%d.%m.%Y')

        # Convert paragraph breaks to Word line breaks (preserve double newlines)
        cover_letter_formatted = cover_letter_text.replace('\n\n', '\n\n')

        replacements = {
            '{{SENDER_NAME}}': self.sender['name'],
            '{{SENDER_PHONE}}': self.sender['phone'],
            '{{SENDER_EMAIL}}': self.sender['email'],
            '{{SENDER_ADDRESS_LINE1}}': self.sender['address_line1'],
            '{{SENDER_ADDRESS_LINE2}}': self.sender['address_line2'],
            '{{SENDER_LINKEDIN}}': self.sender['linkedin'],
            '{{COMPANY_NAME}}': job_data.get('company_name', 'Company'),
            '{{COMPANY_ADDRESS}}': job_data.get('company_address', ''),
            '{{COMPANY_ADDRESS_LINE1}}': job_data.get('company_address_line1', ''),
            '{{COMPANY_ADDRESS_LINE2}}': job_data.get('company_address_line2', ''),
            '{{COMPANY_LOCATION}}': job_data.get('location', ''),
            '{{JOB_TITLE}}': job_data.get('job_title_clean') or job_data.get('job_title', 'Position'),
            '{{DATE}}': today,
            '{{COVER_LETTER_SALUTATION}}': job_data.get('cover_letter_salutation', ''),
            '{{COVER_LETTER_BODY}}': cover_letter_formatted,
            '{{COVER_LETTER_VALEDICTION}}': job_data.get('cover_letter_valediction', ''),
        }

        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # Save document
        try:
            # Normalize the path to use proper OS separators
            output_path = os.path.normpath(output_path)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
        except Exception as e:
            logger.error("Failed to save DOCX %s: %s", output_path, e)
            raise DocumentError(f"Failed to save DOCX: {output_path}") from e

        logger.info("Word document generated: %s", output_path)
        return output_path
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, Any]) -> None:
        """Replace placeholders in a paragraph while preserving formatting"""
        
        # First try: simple replacement within runs
        for run in paragraph.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
        
        # Second try: handle placeholders split across runs
        full_text = paragraph.text
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                # Check if it was already replaced
                already_replaced = any(str(value) in run.text for run in paragraph.runs)
                if not already_replaced:
                    # Placeholder is split - need special handling
                    self._replace_split_placeholder(paragraph, placeholder, str(value))
    
    def _replace_split_placeholder(self, paragraph, placeholder: str, value: str) -> None:
        """Handle placeholders that are split across multiple runs"""
        
        # Build the full text with run boundaries marked
        run_texts = [run.text for run in paragraph.runs]
        full_text = ''.join(run_texts)
        
        if placeholder not in full_text:
            return
        
        # Find placeholder position
        start_idx = full_text.index(placeholder)
        end_idx = start_idx + len(placeholder)
        
        # Find which runs contain the placeholder
        current_pos = 0
        affected_runs = []
        
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if run_end > start_idx and run_start < end_idx:
                # This run overlaps with placeholder
                overlap_start = max(0, start_idx - run_start)
                overlap_end = min(len(run.text), end_idx - run_start)
                affected_runs.append({
                    'index': i,
                    'run': run,
                    'overlap_start': overlap_start,
                    'overlap_end': overlap_end,
                    'run_start': run_start,
                    'run_end': run_end
                })
            
            current_pos = run_end
        
        if not affected_runs:
            return
        
        # Replace text in affected runs
        for i, info in enumerate(affected_runs):
            run = info['run']
            overlap_start = info['overlap_start']
            overlap_end = info['overlap_end']
            
            if i == 0:
                # First run: replace from overlap_start
                before = run.text[:overlap_start]
                after = run.text[overlap_end:]
                run.text = before + value + after
            else:
                # Subsequent runs: remove the placeholder part
                before = run.text[:overlap_start]
                after = run.text[overlap_end:]
                run.text = before + after
    
    def _generate_basic_docx(
        self,
        cover_letter_text: str,
        job_data: Dict[str, Any],
        output_path: str,
        language: str
    ) -> str:
        """
        Generate a basic Word document without template
        Fallback if template doesn't exist
        """
        logger = get_logger(__name__)
        logger.info("Creating basic Word document (no template)...")

        doc = Document()

        # Set narrow margins (2.5cm)
        sections = doc.sections
        for section in sections:
            section.top_margin = 914400  # 2.5cm in EMUs
            section.bottom_margin = 914400
            section.left_margin = 914400
            section.right_margin = 914400

        today = datetime.now().strftime('%d.%m.%Y')
        company_name = job_data.get('company_name', 'Company')
        job_title = job_data.get('job_title', 'Position')

        # Add sender info (right-aligned)
        p = doc.add_paragraph()
        p.alignment = 2  # Right align
        p.add_run(f"{self.sender['name']}\n").bold = True
        p.add_run(f"{self.sender['phone']}\n")
        p.add_run(f"{self.sender['email']}\n")
        p.add_run(f"{self.sender['address_line1']}\n")
        p.add_run(f"{self.sender['address_line2']}\n")
        p.add_run(f"{self.sender['linkedin']}")

        doc.add_paragraph()  # Spacer

        # Company address (left-aligned)
        p = doc.add_paragraph()
        p.add_run(f"{company_name}\n").bold = True
        if job_data.get('company_address'):
            p.add_run(f"{job_data['company_address']}\n")
        if job_data.get('location'):
            p.add_run(f"{job_data['location']}")

        doc.add_paragraph()  # Spacer

        # Date (right-aligned)
        p = doc.add_paragraph(today)
        p.alignment = 2  # Right align

        doc.add_paragraph()  # Spacer

        # Subject
        subject = f"Bewerbung als {job_title}" if language == 'german' else f"Application for {job_title}"
        p = doc.add_paragraph()
        p.add_run(subject).bold = True

        doc.add_paragraph()  # Spacer

        # Greeting
        greeting = f"Hallo liebes {company_name}-Team," if language == 'german' else "Dear Hiring Team,"
        doc.add_paragraph(greeting)

        # Cover letter body
        paragraphs = cover_letter_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        doc.add_paragraph()  # Spacer

        # Closing
        closing = "Mit freundlichen Grüßen," if language == 'german' else "Best regards,"
        doc.add_paragraph(closing)
        doc.add_paragraph()  # Space for signature

        # Name
        p = doc.add_paragraph()
        p.add_run(self.sender['name']).bold = True

        doc.add_paragraph()  # Spacer

        # Attachments
        attachments = "Anlagen: Deckblatt · Lebenslauf · Zeugnisse" if language == 'german' else "Attachments: Cover Page · Resume · References"
        p = doc.add_paragraph(attachments)
        p.runs[0].font.italic = True

        # Save
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
        except Exception as e:
            logger.error("Failed to save basic DOCX %s: %s", output_path, e)
            raise DocumentError(f"Failed to save DOCX: {output_path}") from e

        logger.info("Basic Word document created: %s", output_path)
        return output_path
    
    def convert_to_pdf(self, docx_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
        """
        Convert .docx to PDF
        Note: Requires Microsoft Word to be installed OR LibreOffice
        
        Args:
            docx_path (str): Path to .docx file
            pdf_path (str): Output PDF path (optional)
            
        Returns:
            str: Path to PDF or None if conversion failed
        """
        
        if pdf_path is None:
            pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Try using docx2pdf (works on Windows with Word installed)
            import docx2pdf
            docx2pdf.convert(docx_path, pdf_path)
            logger = get_logger(__name__)
            logger.info("Converted to PDF: %s", pdf_path)
            return pdf_path
        except ImportError:
            logger = get_logger(__name__)
            logger.warning("docx2pdf not installed. Install with: pip install docx2pdf")
            logger.warning("Or manually export to PDF from Word")
            return None
        except Exception as e:
            logger = get_logger(__name__)
            logger.warning("Could not convert to PDF: %s", e)
            logger.warning("You can manually export to PDF from Word")
            return None


# Test function
if __name__ == "__main__":
    test_cover_letter = """Die STRABAG AG beeindruckt mich durch ihren innovativen Ansatz und das Engagement für nachhaltiges Bauen. Als erfahrener Produktmanager mit umfangreicher Expertise in der digitalen Transformation sehe ich großes Potenzial, zur strategischen Weiterentwicklung Ihrer Digital Workplace Lösungen beizutragen.

In meiner letzten Position bei von Rundstedt habe ich eine umfassende strategische Vision für ein digitales B2B-Portfolio entwickelt, das Marktchancen von über 20 Millionen Euro umfasst. Ich leitete ein cross-funktionales Team von zehn Mitgliedern und verantwortete die Finanzierung von über 200.000 Euro.

Meine analytischen Fähigkeiten sowie meine Erfahrung in der Roadmap-Planung und Priorisierung von Features machen mich zu einem geeigneten Kandidaten für die Position. Ich bin hochmotiviert, mein Know-how bei STRABAG einzubringen und gemeinsam an der Zukunft des Bauens zu arbeiten."""
    
    test_job_data = {
        'company_name': 'STRABAG AG',
        'company_address': 'Albstadtweg 10, 70567 Stuttgart',
        'job_title': 'Product Manager Digital Workplace (m/w/d)',
        'location': 'Stuttgart'
    }
    
    print("=" * 80)
    print("WORD TEMPLATE GENERATOR TEST")
    print("=" * 80)
    
    generator = WordCoverLetterGenerator()
    output_file = "output/cover_letters/test_cover_letter.docx"
    
    docx_path = generator.generate_from_template(
        test_cover_letter,
        test_job_data,
        output_file,
        language='german'
    )
    
    print(f"\n✓ Test complete!")
    print(f"✓ Word document: {docx_path}")
    print(f"\nNext steps:")
    print(f"1. Open {docx_path} in Word")
    print(f"2. Adjust formatting/styling as desired")
    print(f"3. Save as your template: data/template_de.docx")
    print(f"4. Create similar template: data/template_en.docx")