"""
Debug script to see how Word stores text runs in the template
"""

from docx import Document

template_path = 'data/template_de.docx'

print("=" * 80)
print("TEMPLATE RUN ANALYSIS")
print("=" * 80)

doc = Document(template_path)

print("\n--- PARAGRAPHS ---")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"\nParagraph {i}: '{para.text[:60]}...'")
        print(f"  Style: {para.style.name}")
        print(f"  Number of runs: {len(para.runs)}")
        
        for j, run in enumerate(para.runs):
            print(f"    Run {j}: '{run.text}'")
            print(f"      Font: {run.font.name}, Size: {run.font.size}")
            print(f"      Bold: {run.font.bold}, Italic: {run.font.italic}")
            print(f"      Color: {run.font.color.rgb if run.font.color and run.font.color.rgb else 'Default'}")

print("\n--- TABLES ---")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, para in enumerate(cell.paragraphs):
                if para.text.strip():
                    print(f"  Cell [{r_idx},{c_idx}] Para {p_idx}: '{para.text[:40]}...'")
                    print(f"    Runs: {len(para.runs)}")
                    for run_idx, run in enumerate(para.runs):
                        print(f"      Run {run_idx}: '{run.text}'")
                        print(f"        Bold: {run.font.bold}, Size: {run.font.size}")

print("\n" + "=" * 80)