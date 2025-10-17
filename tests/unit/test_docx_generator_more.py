import os
import types

import pytest

from src.docx_generator import WordCoverLetterGenerator
from src.utils.errors import DocumentError


def test_replace_split_placeholder_in_paragraph(tmp_path):
    gen = WordCoverLetterGenerator()

    # Create a minimal docx in-memory by generating a basic file, then open and modify
    out = tmp_path / "doc.docx"
    result = gen._generate_basic_docx("Hello", {"company_name": "Acme", "job_title": "Dev"}, str(out), "english")

    from docx import Document
    doc = Document(result)

    # Add a paragraph with split placeholder across runs
    p = doc.add_paragraph()
    run1 = p.add_run("Before ")
    run2 = p.add_run("{{COVER_LETTER_")
    run3 = p.add_run("BODY}} after")

    # Save then run replacement
    doc.save(result)

    # Now reopen via generate_from_template path using a fake template path by temporarily pointing to our file
    gen.template_en = str(out)
    replaced = gen.generate_from_template("BodyText", {"company_name": "Acme", "job_title": "Dev"}, str(out), "english")

    # Reopen and check the paragraph contains replaced text
    doc2 = Document(replaced)
    texts = [p.text for p in doc2.paragraphs if "BodyText" in p.text]
    assert any("BodyText" in t for t in texts)


def test_generate_from_template_handles_table_placeholders(tmp_path):
    gen = WordCoverLetterGenerator()
    out = tmp_path / "doc2.docx"

    from docx import Document
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "{{COMPANY_NAME}}"
    table.cell(0, 1).text = "{{JOB_TITLE}}"
    doc.save(out)

    gen.template_en = str(out)
    res = gen.generate_from_template("Body", {"company_name": "Acme", "job_title": "Dev"}, str(out), "english")

    doc2 = Document(res)
    texts = [cell.text for row in doc2.tables[0].rows for cell in row.cells]
    assert "Acme" in texts[0]
    assert "Dev" in texts[1]


def test_generate_from_template_raises_on_save_error(monkeypatch, tmp_path):
    gen = WordCoverLetterGenerator()
    out = tmp_path / "doc3.docx"

    # Create a valid template to load
    from docx import Document
    doc = Document()
    doc.add_paragraph("{{COVER_LETTER_BODY}}")
    doc.save(out)
    gen.template_en = str(out)

    # Monkeypatch doc.save to throw
    import docx
    original_save = docx.document.Document.save

    def boom(self, path):
        raise OSError("disk full")

    monkeypatch.setattr(docx.document.Document, "save", boom)

    with pytest.raises(DocumentError):
        gen.generate_from_template("Body", {"company_name": "Acme", "job_title": "Dev"}, str(out), "english")

    # Restore not strictly necessary in test, but good hygiene
    monkeypatch.setattr(docx.document.Document, "save", original_save)
